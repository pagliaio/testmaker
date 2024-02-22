import random
import sys
import tkinter as tk
from tkinter import *
from tkinter import ttk
from tkinter import messagebox
from docx import Document
import re
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import ctypes
from tkinter import filedialog

from docx import Document

from tkinter import simpledialog
from tkinter import ttk


# This functions allows the program to write the Title and the Header in bold
def set_run_bold(run):
    run.font.bold = True
    r_elem = run._element
    b = OxmlElement('w:b')
    r_elem.insert(0, b)

# This function adds the header to the test
def add_header(document, title):
    p = document.add_paragraph(style='Normal')
    p.add_run(title).font.bold = True
    p.add_run(f"\n" + student_name_string.get() + " ___________________________________    "+ student_class_string.get() + " ________  " + date_string.get() + " ________________\n")

def add_page_break(document):
    document.add_page_break()

def check_text_not_empty(text):
    return bool(text.strip()) if isinstance(text, str) else bool(text)
    
def check_list_not_empty(text_list):
    return any(text.strip() for text in text_list)

def add_fitb(document, words):
    print("Received triplets:")  # Print a header for clarity
    for word in words:
        print(word)  # Print each triplet
    print()  # Print an empty line for clarity
    
    table = document.add_table(rows=1, cols=1)
    table.autofit = True  # Autofit table to content

    if words:
        row_cells = table.rows[0].cells
        # Join all words into a single string separated by commas, adding bold formatting
        row_cells[0].paragraphs[0].add_run(", ".join(words)).bold = True  
        row_cells[0].paragraphs[0].alignment = 1  # Center-align the text within the cell


    

def add_flexible_exercise(document, exercise, exercise_index):
    p = document.add_paragraph()
    
    run = p.add_run(exercise_name_string.get() + f" {exercise_index} - {exercise['task_description']}")
    set_run_bold(run)
    
    # Add reading test if present
    readtext = exercise['hastext']
    if readtext:
        document.add_paragraph(f"{exercise['reading_text']}", style='Normal')

    sentences = []  # Collect all sentences into a single list
    words = []  # Collect all words into a single list
    mcoptions = []  # Collect all options into a single list

    for triplet in exercise['sentence_fitb_mco']:
        print("Processing triplet:", triplet)  # Debug print
        if triplet[0]:  # Check if the sentence element of the triplet is not empty
            sentences.append(triplet[0])  # Add the sentence to the list of sentences
        
        if triplet[1]:  # Check if the word element of the triplet is not empty
            words.append(triplet[1])  # Add the word to the list of words
        
        if len(triplet) >= 3 and triplet[2]:  # Check if the options element of the triplet is not empty
            mcoptions.append(triplet[2])  # Add the options to the list of options

    print("Sentences:", sentences)
    print("Words:", words)
    print("Multiple Choice Options:", mcoptions)

    underscore_line = "_" * 100
    
    list_style = exercise['list_style']
    print("Style:", list_style)
 # Set sentences_per_student equal to the total number of sentences if it's 0
    sentences_per_student = exercise['sentences_per_student'] if exercise['sentences_per_student'] > 0 else len(exercise['sentence_fitb_mco'][0])
  


    # Function to shuffle options if mc_to_be_scrambled is 1
    def shuffle_options(options):
        if options is None:  # Check if options is None
            return None
        if exercise['mc_to_be_scrambled'] == 1:
            options = list(options)  # Convert options to a list
            random.shuffle(options)
            options = tuple(options)  # Convert options back to a tuple
        return options
    
    # Check if there are sentences to print
    if exercise['hassentences'] == 1:
        # la generazione di frasi con linee funziona
        if exercise['underscore_lines'] != 0:
            if exercise['to_be_scrambled'] == 1:
                random.shuffle(sentences)
            if exercise['avoid_repetitions'] == 1:
                random.shuffle(sentences)    
                selected_sentences = sentences[:sentences_per_student]
                for sentence in selected_sentences:
                    document.add_paragraph(sentence, style=list_style)
                    for _ in range(exercise['underscore_lines']):
                        document.add_paragraph("_" * 100, style='Normal')
            else:
                for sentence in sentences:
                    document.add_paragraph(sentence, style=list_style)
                    for _ in range(exercise['underscore_lines']):
                        document.add_paragraph("_" * 100, style='Normal')
                    
                    
        # Check for multiple-choice questions if multiple_choice_option_number > 0
        elif exercise['multiple_choice_option_number'] > 0:
            # Check if multiple_choice_option_number is 1 and mcoptions list is empty
            if exercise['multiple_choice_option_number'] == 1 and not mcoptions:
                shuffled_sentences = list(set(sentences))  # Get unique sentences
                if exercise['to_be_scrambled'] == 1:
                        random.shuffle(shuffled_sentences)
                if exercise['avoid_repetitions'] == 1:    
                    selected_sentences = shuffled_sentences[:sentences_per_student]
                    for index, sentence in enumerate(selected_sentences, start=1):
                        document.add_paragraph(f"{index}) {sentence}", style= list_style)
                        document.add_paragraph("\n\t\tT)\tF)", style='Normal')
                else:
                    for index, sentence in enumerate(sentences, start=1):
                        document.add_paragraph(f"{index}) {sentence}", style= list_style)
                        document.add_paragraph("\n\t\tT)\tF)", style='Normal')
            else:
                if exercise['avoid_repetitions'] ==1:
                    indices = list(range(len(sentences)))
                    random.shuffle(indices)
                    selected_indices = indices[:sentences_per_student]

                    for index in selected_indices:
                        sentence = sentences[index]
                        options = mcoptions[index]

                        # Shuffle options if specified
                        if exercise.get('mc_to_be_scrambled', 0) == 1:
                            options = list(options)
                            random.shuffle(options)
                            options = tuple(options)

                        document.add_paragraph(sentence, style= list_style)  # Print the sentence without index

                        option_string = ""
                        for i, option in enumerate(options, start=65):  # Start from ASCII 'A'
                            option_string += f"{chr(i)}. {option}\t\t"
                        document.add_paragraph(option_string.rstrip('\t\t'), style='Normal')
                else:
            # Shuffle sentences if to_be_scrambled is enabled
                    if exercise['to_be_scrambled'] == 1: 
                        indices = list(range(len(sentences)))
                        random.shuffle(indices)
                        for index in indices:
                            sentence = sentences[index]
                            options = shuffle_options(mcoptions[index])  # Shuffle options
                            document.add_paragraph(sentence, style= list_style)
                            option_string = ""
                            for i, option in enumerate(options, start=65):  # Start from ASCII 'A'
                                option_string += f"{chr(i)}. {option}\t\t"
                            document.add_paragraph(option_string.rstrip('\t\t'), style='Normal')
                    else:
                        for sentence, options in zip(sentences, mcoptions):
                            options = shuffle_options(options)  # Shuffle options
                            document.add_paragraph(sentence, style=list_style)
                            option_string = ""
                            for i, option in enumerate(options, start=65):  # Start from ASCII 'A'
                                option_string += f"{chr(i)}. {option}\t\t"
                            document.add_paragraph(option_string.rstrip('\t\t'), style='Normal')
        else:
            # la generazione di frasi senza multiple choice o linee funziona
            if exercise['hasfitb']:
                print("AVOID REPETITIONS:")
                print(exercise['avoid_repetitions'])
                if exercise['avoid_repetitions'] == 1: 
                    
                    indices = list(range(len(sentences)))
                    random.shuffle(indices)
                    selected_indices = indices[:sentences_per_student]
                    
                    fitbwords = []
                    for index in selected_indices:
                            fitbwords.append(words[index])
                            
                    fitbsentences = []
                    add_fitb(document, fitbwords)        
                    random.shuffle(fitbsentences)
                    for index in selected_indices:
                            fitbsentences.append(sentences[index]) 
                    random.shuffle(fitbsentences)        
                    for sentence in fitbsentences:
                        document.add_paragraph(sentence, style=list_style)                
                else:
                    if exercise['to_be_scrambled'] == 1:
                        random.shuffle(sentences)
                    add_fitb(document, words)
                    for sentence in sentences:
                            document.add_paragraph(sentence, style=list_style)
            else:
                if exercise['to_be_scrambled'] == 1:
                        random.shuffle(sentences)
                if exercise['avoid_repetitions'] == 1:
                        random.shuffle(sentences)    
                        selected_sentences = sentences[:sentences_per_student]
                        for sentence in selected_sentences:
                            document.add_paragraph(sentence, style=list_style)

                else:
                        for sentence in sentences:
                            document.add_paragraph(sentence, style=list_style)

    else:
        # Handle the case when there are no sentences
        if exercise['hasfitb']:
            add_fitb(document, words)
        elif exercise['multiple_choice_option_number'] > 1:
                        for options in mcoptions:
                            option_string = ""
                            options = shuffle_options(options)  # Shuffle options

                            option_string = ""
                            for i, option in enumerate(options, start=65):  # Start from ASCII 'A'
                                option_string += f"{chr(i)}. {option}\t\t"
                            document.add_paragraph(option_string.rstrip('\t\t'), style=list_style)



        elif exercise['underscore_lines'] != 0:
            # Print underscore lines
            for _ in range(exercise['underscore_lines']):
                document.add_paragraph("_" * 100, style='Normal')

    document.add_paragraph()


# this function is called when the Generate button is pressed.
def generate_test(title, output_filename, students, exercise_entries):
    document = Document()
    
    for student in range(1, students + 1):
        add_header(document, title)
        for exercise_index, exercise_entry in enumerate(exercise_entries, start=1):
            add_flexible_exercise(document, exercise_entry, exercise_index)
        print("Exercise index", exercise_index)
        print("len", len(exercise_entries))
        print("oddp ages ", odd_pages_checkbox_value.get())
        if exercise_index == len(exercise_entries):
                add_page_break(document)
                if odd_pages_checkbox_value.get() == 1:
                    add_page_break(document)
    # Get the default directory
    default_dir = filedialog.askdirectory()

    # Combine directory and output filename
    save_path = f"{default_dir}/{output_filename}"

    if save_path:
        document.save(save_path)

        # Custom messagebox with "Open file location" button
        result = messagebox.askquestion("Info", f"Output saved to {save_path}. Would you like to open the file location?",
                                         icon='info')
        if result == 'yes':
            import os
            os.startfile(os.path.dirname(os.path.abspath(save_path)))



# questa funzione accoppia (attripla) frase, parola per riempimento, e risposte del multiple choice. in questo modo uso una cosa per storare vari tipi di esercizio
def generate_sentence_word_mco(sentences, vocabulary_words, list_of_lists_of_mcoptions):

    sentence_word_mco = []

    while True:
        sentence = sentences.pop(0).strip() if sentences else None
        word = vocabulary_words.pop(0).strip() if vocabulary_words else None  # Remove leading and trailing whitespace from word
        options = list_of_lists_of_mcoptions.pop(0) if list_of_lists_of_mcoptions else None

        # Check if all three elements are None
        if sentence is None and word is None and options is None:
            break  # If all three are None, break out of the loop
        
        # Append the triplet to sentence_word_mco regardless of whether any element is empty
        sentence_word_mco.append((sentence, word, options))
        sentence = sentence or ""
        word = word or ""
        options = options or []

    print(sentence_word_mco)
    return sentence_word_mco



# This function adds to the document the current exercise. It then resets the fields in the form.
def add_to_test():
    task_description = task_description_entry.get("1.0", "end-1c")
    if not task_description:
        messagebox.showwarning("Warning", "Please fill in the task description.")
        return
    
    # Prendi le frasi tutte in blocco, senza splittare, e guarda se ce ne sono
    sentencestext = sentences_entry.get("1.0", "end-1c").split("\n")
    sentences = sentences_entry.get("1.0", "end-1c")
    hassentences = check_text_not_empty(sentences)    
    
    # Prendi le parole per il fitb tutte in blocco, senza splittare, e guarda se ce ne sono
    vocabulary_words = [word.strip() for word in fitb_entry.get("1.0", "end-1c").split(",")]
    hasfitb = check_list_not_empty(vocabulary_words)
    print(hasfitb)  
    print(vocabulary_words)   
    multiple_choice_answers_text = multiple_choice_answers_entry.get("1.0", "end-1c")
   # Prendi le risposte per il multiple choice tutte in blocco, splitta per riga e poi per virgola
    mcoptions = [[item.strip() for item in line.split(",")] for line in multiple_choice_answers_text.split("\n") if line.strip()]
    
    # Controlla se ci sono da mettere multiple options
    multiple_choice_option_number = multiple_choice_options_value.get()
    
    #appaia sentences, fitb words e risposte per multiple options in un'unica lista
    sentence_fitb_mco = generate_sentence_word_mco(sentencestext, vocabulary_words, mcoptions)
    
    
    # se ci sono delle linee da aggiungere per far scrivere gli studenti (verranno messe tutte in una fila nel caso in cui siano assenti le Sentences/Domande, perché si darà per buono
    # che in quel caso le linee serviranno per un esercizio di writing. Viceversa, nel caso delle frasi presenti, queste saranno intervallate dal numero di righe specificate da questo valore)
    underscore_lines = int(underscore_lines_var.get())
    underscore_string = underscore_lines_var.get()
    if not check_text_not_empty(underscore_string):
        underscore_lines = 0   
        

        
    # controlla se ci c'è un testo da leggere dopo la consegna
    reading_text = reading_test_entry.get("1.0", "end-1c")
    readtext = reading_test_entry.get("1.0", "end-1c")
    hastext = check_text_not_empty(readtext)    
    
    
    # il numero di frasi per studente dal pool totale viene messo a 0 (che equivale a "tutti hanno lo stesso numero di frasi, il pool intero di norma" anche se il campo è vuoto)
    sentences_per_student = int(sentences_per_student_entry.get())
    sentences_per_student_string = sentences_per_student_entry.get()
    if not check_text_not_empty(sentences_per_student_string):
        sentences_per_student = 0
     
    # bisogna mescolare in modo random le frasi? le frasi sono da scrivere su due colonne? bisogna evitare ripetizioni nel caso in cui 
    # ciasun studente abbia solo una parte del pool totale di frasi?
    if sentences_per_student == 0:
        avoid_repetitions = 0
    else:
        avoid_repetitions = 1
        
        
        
    print(sentences_per_student)
    print(avoid_repetitions)
    if style_value.get() == 0:
        list_style = "Normal"
    else:
        list_style = "ListBullet"
    
    to_be_scrambled = scramble_checkbox_value.get()

 
    mc_to_be_scrambled = scramble_mc_checkbox_value.get()

    exercise_entry = {'task_description': task_description, 'underscore_lines': underscore_lines, 'reading_text': reading_text, 'sentence_fitb_mco':sentence_fitb_mco, 'list_style': list_style,
                      'hasfitb' : hasfitb, 'to_be_scrambled':to_be_scrambled, 'mc_to_be_scrambled': mc_to_be_scrambled, 'hassentences': hassentences, 'hastext': hastext,
                      'sentences_per_student' : sentences_per_student, 'avoid_repetitions' : avoid_repetitions, 'multiple_choice_option_number' : multiple_choice_option_number}
    print("Has fitb?")
    print(hasfitb)
    print(vocabulary_words)
    # resetta l'interfaccia
    exercise_entries.append(exercise_entry)
    task_description_entry.delete("1.0", tk.END)
    sentences_entry.delete("1.0", tk.END)
    fitb_entry.delete("1.0", tk.END)
    reading_test_entry.delete("1.0", tk.END)
    underscore_lines_var.set(0)
    multiple_choice_answers_entry.delete("1.0", tk.END)
    multiple_choice_options_value.set(0)    
    scramble_mc_checkbox_value.set(0)

    sentences_per_student_value.set(0)  




# Create GUI
root = tk.Tk()
root.title("Albe's Test Maker")
root.configure(background='darkgray')

# Get the screen width and height
user32 = ctypes.windll.user32
screen_width = user32.GetSystemMetrics(0)
screen_height = user32.GetSystemMetrics(1)
root.geometry(f"{screen_width}x{screen_height}")
root.state('zoomed')  # For Windows and Linux

create_frame = tk.Frame (root, width=400, height=200, bg='lightgrey')
create_frame.grid(row=0, column=0, padx=10, pady=5, sticky="n") 

title_label = tk.Label (create_frame, text='Title',  width=20, height=1, anchor="e")
title_label.grid(row=0, column=0, padx=5, pady=5, sticky='ew')
title_entry = tk.Entry(create_frame)
title_entry.grid(row=0, column=1, padx=5, pady=5)

filename_label = tk.Label (create_frame, text='Output file name', width=20, height=1, anchor="e")
filename_label.grid(row=1, column=0, padx=5, pady=5, sticky='ew')
filename_entry = tk.Entry(create_frame)
filename_entry.grid(row=1, column=1, padx=5, pady=5)

students_label = tk.Label (create_frame, text='Number of students', width=20, height=1, anchor="e")
students_label.grid(row=2, column=0, padx=5, pady=5, sticky='ew')
students_entry = tk.Entry(create_frame)
students_entry.grid(row=2, column=1, padx=5, pady=5)

def reset_test():
    result = messagebox.askquestion("Warning", "Are you sure you want to reset the test?")
    if result == "yes":
        global exercise_entries
        exercise_entries=[]
        task_description_entry.delete("1.0", tk.END)
        sentences_entry.delete("1.0", tk.END)
        fitb_entry.delete("1.0", tk.END)
        reading_test_entry.delete("1.0", tk.END)
        underscore_lines_var.set(0)
        multiple_choice_answers_entry.delete("1.0", tk.END)
        multiple_choice_options_value.set(0)    
        scramble_mc_checkbox_value.set(0)
        title_entry.delete(0, tk.END)
        filename_entry.delete(0, tk.END)
        students_entry.delete(0, tk.END)
        sentences_per_student_value.set(0)  
        print("Test reset")  # Placeholder action, replace with actual reset logic
    else:
        print("Reset canceled")
    
reset_button = tk.Button(create_frame, text="Reset", command=reset_test)
reset_button.grid(row=4, column=1, padx=5, pady=5)



generate_button = tk.Button(create_frame, text="Generate", bg="crimson", command=lambda: generate_test(title_entry.get(), filename_entry.get() + ".docx",
                                                           int(students_entry.get()), exercise_entries))
generate_button.grid(row=3, column=1, padx=5, pady=5)


#reset_button = tk.Button(create_frame, text="Reset", command=lambda: reset_test(title_entry.get(), filename_entry.get() + ".docx",
#                                                           int(students_entry.get()), exercise_entries))
#reset_button.grid(row=3, column=1, padx=5, pady=5)

scramble_checkbox_value = tk.IntVar()
scramble_checkbox_value.set(1)

scramble_checkbox = tk.Checkbutton(create_frame, text="Scramble", variable=scramble_checkbox_value, anchor="e", bg="lightcoral")
scramble_checkbox.grid(row=3, column=0, padx=5, pady=5, sticky='ew')



odd_pages_checkbox_value = tk.IntVar()
odd_pages_checkbox_value.set(0)

odd_pages_checkbox = tk.Checkbutton(create_frame, text="Odd pages (3 or more)", variable=odd_pages_checkbox_value, anchor="e")
odd_pages_checkbox.grid(row=4, column=0, padx=5, pady=5, sticky='ew')

sentences_per_student_value = tk.IntVar()
sentences_per_student_label = tk.Label (create_frame, text='Sentences per student\n(Leave empty to include all)', width=25, height=2, anchor="e")
sentences_per_student_label.grid(row=5, column=0, padx=5, pady=5, sticky='ew')
sentences_per_student_entry = tk.Entry(create_frame, width=10, textvariable=sentences_per_student_value)
sentences_per_student_entry.grid(row=5, column=1, padx=5, pady=5)

style_value = tk.IntVar()
style_value.set(0)
style_repetitions_radio_button1 = tk.Radiobutton(create_frame, text="Bullet points", variable=style_value, value=1, bg="yellow")
style_repetitions_radio_button1.grid(row=6, column=0, padx=5, pady=5, sticky='ew')

style_repetitions_radio_button2 = tk.Radiobutton(create_frame, text="Normal", variable=style_value, value=0, bg="yellow")
style_repetitions_radio_button2.grid(row=6, column=1, padx=5, pady=5, sticky='ew')



underscore_lines_label = tk.Label (create_frame, text='Writing lines', width=20, height=1, anchor="e", bg="lightblue")
underscore_lines_label.grid(row=7, column=0, padx=5, pady=5, sticky='ew')
underscore_lines_var = tk.StringVar(create_frame)
underscore_lines_var  = tk.IntVar()
underscore_lines_entry = tk.Entry(create_frame, textvariable=underscore_lines_var, bg="lightblue")
underscore_lines_entry.grid(row=7, column=1, padx=5, pady=5, sticky='ew')


student_name_string_label = tk.Label (create_frame, text='Student:', width=20, height=1, anchor="e")
student_name_string_label.grid(row=8, column=0, padx=5, pady=5, sticky='ew')
student_name_string = tk.StringVar(create_frame)
student_name_string.set("Student")
student_name_string__entry = tk.Entry(create_frame, textvariable=student_name_string)
student_name_string__entry.grid(row=8, column=1, padx=5, pady=5, sticky='ew')

student_class_string_label = tk.Label (create_frame, text='Class:', width=20, height=1, anchor="e")
student_class_string_label.grid(row=9, column=0, padx=5, pady=5, sticky='ew')
student_class_string = tk.StringVar(create_frame)
student_class_string.set("Class")
student_class_string__entry = tk.Entry(create_frame, textvariable=student_class_string)
student_class_string__entry.grid(row=9, column=1, padx=5, pady=5, sticky='ew')

date_string_label = tk.Label (create_frame, text='Date:', width=20, height=1, anchor="e")
date_string_label.grid(row=10, column=0, padx=5, pady=5, sticky='ew')
date_string = tk.StringVar(create_frame)
date_string.set("Date")
date_string__entry = tk.Entry(create_frame, textvariable=date_string)
date_string__entry.grid(row=10, column=1, padx=5, pady=5, sticky='ew')

exercise_name_string_label = tk.Label (create_frame, text='Exercise:', width=20, height=1, anchor="e")
exercise_name_string_label.grid(row=11, column=0, padx=5, pady=5, sticky='ew')
exercise_name_string = tk.StringVar(create_frame)
exercise_name_string.set("Exercise")
exercise_name_string__entry = tk.Entry(create_frame, textvariable=exercise_name_string)
exercise_name_string__entry.grid(row=11, column=1, padx=5, pady=5, sticky='ew')



# EXERCISE FRAME
exercise_frame = tk.Frame (root, width=500, height=500, bg='lightgray')
exercise_frame.grid(row=0, column=2, padx=5, pady=5, sticky="ew") 

# Task Description
task_description_label = tk.Label (exercise_frame, text='Task description', width=18, height=1, anchor="center")
task_description_label.grid(row=0, column=0, padx=5, pady=5, sticky='ew')
task_description_entry = tk.Text(exercise_frame,wrap="none", height=1, width=20)
task_description_entry.grid(row=0, column=1, padx=5, pady=5, columnspan=2, sticky='ew')

# Sentences
sentences_label = tk.Label (exercise_frame, text='Sentences/Closed questions:', width=25, height=1, anchor="center", bg="burlywood")
sentences_label.grid(row=1, column=0, padx=5, pady=5, sticky='sw')
sentences_entry = tk.Text(exercise_frame, height=30, width=80, bg="burlywood")
sentences_entry.grid(row=2, column=0, padx=5, pady=5, columnspan=2, sticky='nw')

# Fill in the blank words
fitb_label = tk.Label (exercise_frame, text='Fill in the blank words:', width=30, height=1, anchor="center", bg='violet')
fitb_label.grid(row=3, column=0, padx=5, pady=5, sticky='ew')
fitb_entry = tk.Text(exercise_frame, height=5, width=30, bg='violet')
fitb_entry.grid(row=4, column=0, padx=5, pady=5, columnspan=2, sticky='ew')

# Reading Test (for Reading exercises)
reading_test_label = tk.Label (exercise_frame, text='Text to read:', width=30, height=1, anchor="center", bg='orange')
reading_test_label.grid(row=3, column=1, padx=5, pady=5, sticky='ew')
reading_test_entry = tk.Text(exercise_frame, height=5, width=30, bg='orange')
reading_test_entry.grid(row=4, column=1, padx=5, pady=5, sticky='ew')

# Multiple choice answers
multiple_choice_answers_label = tk.Label (exercise_frame, text='Multiple choice answers:', width=30, height=1, anchor="center", bg='lawngreen')
multiple_choice_answers_label.grid(row=5, column=1, padx=5, pady=5, sticky='ew')
multiple_choice_answers_entry = tk.Text(exercise_frame, height=10, width=30, bg='lawngreen')
multiple_choice_answers_entry.grid(row=6, column=1, padx=5, pady=5, sticky='ew')


multiple_choice_options_frame = tk.Frame (exercise_frame, width=20, height=100, bg='lawngreen')
multiple_choice_options_frame.grid(row=6, column=0, padx=5, pady=5, sticky="ew") 
multiple_choice_options_label = tk.Label (exercise_frame, text='Multiple choice options:', width=30, height=1, anchor="center", bg='lawngreen')
multiple_choice_options_label.grid(row=5, column=0, padx=5, pady=5, sticky='ew')
multiple_choice_options_value = tk.IntVar()
multiple_choice_options_value.set(0)
multiple_choice_options_radio_button1 = tk.Radiobutton(multiple_choice_options_frame, text="None", variable=multiple_choice_options_value, value=0, bg="lawngreen")
multiple_choice_options_radio_button1.grid(row=0, column=0, padx=5, pady=5, sticky='ew')

multiple_choice_options_radio_button2 = tk.Radiobutton(multiple_choice_options_frame, text="True/False", variable=multiple_choice_options_value, value=1, bg="lawngreen")
multiple_choice_options_radio_button2.grid(row=1, column=0, padx=5, pady=5, sticky='ew')

multiple_choice_options_radio_button2 = tk.Radiobutton(multiple_choice_options_frame, text="2 choices", variable=multiple_choice_options_value, value=2, bg="lawngreen")
multiple_choice_options_radio_button2.grid(row=2, column=0, padx=5, pady=5, sticky='ew')

multiple_choice_options_radio_button2 = tk.Radiobutton(multiple_choice_options_frame, text="3 choices", variable=multiple_choice_options_value, value=3, bg="lawngreen")
multiple_choice_options_radio_button2.grid(row=0, column=1, padx=5, pady=5, sticky='ew')

multiple_choice_options_radio_button2 = tk.Radiobutton(multiple_choice_options_frame, text="4 choices", variable=multiple_choice_options_value, value=4, bg="lawngreen")
multiple_choice_options_radio_button2.grid(row=1, column=1, padx=5, pady=5, sticky='ew')

multiple_choice_options_radio_button2 = tk.Radiobutton(multiple_choice_options_frame, text="5 choices", variable=multiple_choice_options_value, value=5, bg="lawngreen")
multiple_choice_options_radio_button2.grid(row=2, column=1, padx=5, pady=5, sticky='ew')

scramble_mc_checkbox_value = tk.IntVar()
scramble_mc_checkbox_value.set(1)

scramble_mc_checkbox = tk.Checkbutton(multiple_choice_options_frame, text="Scramble", variable=scramble_mc_checkbox_value, anchor="e", bg="lawngreen")
scramble_mc_checkbox.grid(row=3, column=1, padx=5, pady=5, sticky='ew')

# CATALOGUE FRAME




def extract_titles(filename):
    titles = []
    with open(filename, 'r') as file:
        for line in file:
            line = line.strip()
            match = re.match(r'\[TITLE\](.*)', line)
            if match:
                titles.append(match.group(1))
    return titles


catalogue_frame = tk.Frame (root, width=300, height=1000, bg='lightyellow')
catalogue_frame.grid(row=0, column=3, padx=10, pady=5, sticky="ne") 



def open_file():
    filename = filedialog.askopenfilename(filetypes=[("Text files", "*.txt")])
    if filename:
        exercise_titles = extract_titles(filename)
        catalogue_menu.config(values=exercise_titles)
        global exercises
        exercises = extract_exercises(filename)

add_button = tk.Button(exercise_frame, text="Add to test", width=20, command=add_to_test, bg="mediumturquoise")
add_button.grid(row=7, column=0, padx=5, pady=5, sticky='ew')



def refresh_combobox():
    # Get the updated exercise titles
    exercise_titles = [exercise.get('TITLE', '') for exercise in exercises]
    # Update the values in the combobox
    catalogue_menu.config(values=exercise_titles)


def add_to_catalogue():
    # Prompt the user to enter the exercise title
    exercise_title = tk.simpledialog.askstring("Add to Catalogue", "Enter the exercise title:")
    if not exercise_title:
        return  # Cancelled by user
    
    # Prompt the user to choose whether to create a new TXT file or add to an existing one
    choice = tk.messagebox.askyesno("Add to Catalogue", "Do you want to create a new TXT file?")
    
    # Open the file dialog to choose the file
    if choice:
        # If creating a new file, ask for the file name
        filename = filedialog.asksaveasfilename(defaultextension=".txt", filetypes=[("Text files", "*.txt")])
        if not filename:
            return  # Cancelled by user
        
        mode = 'w'  # Write mode for a new file
    else:
        # If adding to an existing file, ask for the file name
        filename = filedialog.askopenfilename(filetypes=[("Text files", "*.txt")])
        if not filename:
            return  # Cancelled by user
        
        mode = 'a'  # Append mode for an existing file
    
    # Open the file in the selected mode
    with open(filename, mode) as file:
        # Write the exercise start tag
        file.write("[EXERCISE_START]\n")
        
        # Write the exercise title
        file.write(f"[TITLE]{exercise_title}\n")
        
        # Write the task description to the file
        file.write(f"[TASK_DESCRIPTION]{task_description_entry.get('1.0', tk.END)}")
        
        # Write the sentences to the file
        sentences = sentences_entry.get('1.0', tk.END).strip().split('\n')
        for sentence in sentences:
            file.write(f"[SENTENCES]{sentence}\n")
        
        # Write the fill in the blank words to the file
        fill_in_the_blank_words = fitb_entry.get('1.0', tk.END).strip().split(', ')
        for word in fill_in_the_blank_words:
            file.write(f"[FILL_IN_THE_BLANK_WORDS]{word}\n")
        
        # Write the multiple choice options to the file
        multiple_choice_options = multiple_choice_answers_entry.get('1.0', tk.END).strip().split('\n')
        for options in multiple_choice_options:
            if options:
                file.write(f"[MULTIPLE_CHOICE_OPTIONS]{options}\n")
        
        # Write the reading text to the file
        file.write(f"[READING_TEXT]{reading_test_entry.get('1.0', tk.END)}")
        
        # Write the multiple choice options value to the file
        file.write(f"[MULTIPLE_CHOICE_OPTIONS_VALUE]{multiple_choice_options_value.get()}\n")
        
        # Write the underscore lines variable to the file
        file.write(f"[UNDERSCORE_LINES_VAR]{underscore_lines_var.get()}\n")
        
        # Write the exercise end tag
        file.write("[EXERCISE_END]\n")
    
    tk.messagebox.showinfo("Add to Catalogue", "Data added successfully!")
    
    # Refresh the combobox
    refresh_combobox()



add_to_txt_button = tk.Button(exercise_frame, text="Add to catalogue", width=20, command=add_to_catalogue)
add_to_txt_button.grid(row=7, column=1, padx=5, pady=5, sticky='ew')

        


browse_button = tk.Button(catalogue_frame, text="Browse File", command=open_file)
browse_button.grid(row=0, column=0, pady=5)


def parse_exercises(filename):
    exercises = []
    current_exercise = {}
    with open(filename, 'r') as file:
        for line in file:
            line = line.strip()
            if line.startswith("[EXERCISE_START]"):
                current_exercise = {}
            elif line.startswith("[EXERCISE_END]"):
                exercises.append(current_exercise)
            else:
                match = re.match(r'\[(.*?)\](.*)', line)
                if match:
                    key, value = match.groups()
                    if key == 'SENTENCES':
                        current_exercise.setdefault(key, []).append(value)
                    elif key == 'FILL_IN_THE_BLANK_WORDS':
                        current_exercise.setdefault(key, []).extend(value.split())
                    elif key == 'MULTIPLE_CHOICE_OPTIONS':
                        current_exercise.setdefault(key, []).append(value.split(','))
                    else:
                        current_exercise[key] = value
    return exercises



def extract_exercises(filename):
    exercises = []
    current_exercise = {}
    with open(filename, 'r') as file:
        for line in file:
            line = line.strip()
            if line.startswith("[EXERCISE_START]"):
                current_exercise = {}
            elif line.startswith("[EXERCISE_END]"):
                exercises.append(current_exercise)
            else:
                match = re.match(r'\[(.*?)\](.*)', line)
                if match:
                    key, value = match.groups()
                    if key == 'SENTENCES':
                        current_exercise.setdefault(key, []).append(value)
                    elif key == 'FILL_IN_THE_BLANK_WORDS':
                        current_exercise.setdefault(key, []).append(value)
                    elif key == 'MULTIPLE_CHOICE_OPTIONS':
                        current_exercise.setdefault(key, []).append(value.split(','))
                    else:
                        current_exercise[key] = value
    return exercises

def display_exercise(event):
    selected_title = catalogue_menu.get()
    for exercise in exercises:
        if exercise.get('TITLE') == selected_title:
            display_data(exercise)
            break
        

def display_data(exercise):
    # Clear previous data in listbox
    exercises_catalogue.delete(0, tk.END)
    
    # Insert exercise data into listbox
    exercises_catalogue.insert(tk.END, f"Title: {exercise.get('TITLE', '')}")
    exercises_catalogue.insert(tk.END, f"Task Description:\n{exercise.get('TASK_DESCRIPTION', '')}\n")
    
    sentences = exercise.get('SENTENCES', [])
    if sentences:
        exercises_catalogue.insert(tk.END, "Sentences:")
        for idx, sentence in enumerate(sentences, start=1):
            exercises_catalogue.insert(tk.END, f"  {idx}. {sentence}")
        exercises_catalogue.insert(tk.END, "")
    
    fill_in_the_blank_words = exercise.get('FILL_IN_THE_BLANK_WORDS', [])
    if fill_in_the_blank_words:
        exercises_catalogue.insert(tk.END, f"Fill in the Blank Words: {', '.join(fill_in_the_blank_words)}\n")
    
    multiple_choice_options = exercise.get('MULTIPLE_CHOICE_OPTIONS', [])
    if multiple_choice_options:
        exercises_catalogue.insert(tk.END, "Multiple Choice Options:")
        for idx, options in enumerate(multiple_choice_options, start=1):
            exercises_catalogue.insert(tk.END, f"  {idx}. {', '.join(options)}")
        exercises_catalogue.insert(tk.END, "")
    
    reading_text = exercise.get('READING_TEXT', '')
    if reading_text:
        exercises_catalogue.insert(tk.END, f"Reading Text:\n{reading_text}\n")

    # Set the value of the global tk.IntVar variable for multiple_choice_options_value
    multiple_choice_options_value.set(exercise.get('MULTIPLE_CHOICE_OPTIONS_VALUE', 0))
    
    # Set the value of the global tk.IntVar variable for underscore_lines_var
    underscore_lines_var.set(exercise.get('UNDERSCORE_LINES_VAR', 0))


catalogue_menu= ttk.Combobox(catalogue_frame, width=60)
catalogue_menu.grid(row=0, column=1, padx=5, pady=5, sticky="ne")
catalogue_menu.bind("<<ComboboxSelected>>", display_exercise)



# Create a Treeview to display exercise data
exercises_catalogue = tk.Listbox(catalogue_frame, height=50, width=80, selectmode=EXTENDED )
exercises_catalogue.grid(row=1, column=0, padx=5, pady=5, sticky=tk.NSEW, columnspan=2)


def clear_entry_widgets():
    # Clear previous data in each entry text widget
    task_description_entry.delete('1.0', tk.END)
    sentences_entry.delete('1.0', tk.END)
    fitb_entry.delete('1.0', tk.END)
    multiple_choice_answers_entry.delete('1.0', tk.END)
    reading_test_entry.delete('1.0', tk.END)

def copy_selected_items():
    # Get the selected exercise title from the combobox
    selected_title = catalogue_menu.get()
    
    # Find the corresponding exercise in the list of exercises
    for exercise in exercises:
        if exercise.get('TITLE') == selected_title:
            # Clear previous data in each entry text widget
            clear_entry_widgets()
            
            # Copy task description to task_description_entry
            task_description_entry.insert(tk.END, exercise.get('TASK_DESCRIPTION', '') + '\n\n')
            
            # Copy sentences to sentences_entry
            sentences = exercise.get('SENTENCES', [])
            for sentence in sentences:
                sentences_entry.insert(tk.END, f"{sentence}\n")
            
            # Copy fill in the blank words to fitb_entry
            fill_in_the_blank_words = exercise.get('FILL_IN_THE_BLANK_WORDS', [])
            if fill_in_the_blank_words:
                fitb_entry.insert(tk.END, ', '.join(fill_in_the_blank_words) + '\n\n')
            
            # Copy multiple choice options to multiple_choice_answers_entry
            multiple_choice_options = exercise.get('MULTIPLE_CHOICE_OPTIONS', [])
            for options in multiple_choice_options:
                multiple_choice_answers_entry.insert(tk.END, ', '.join(options) + '\n')
            
            # Copy reading text to reading_text_entry
            reading_text = exercise.get('READING_TEXT', '')
            if reading_text:
                reading_test_entry.insert(tk.END, f'{reading_text}\n\n')

            # Set the value of the global tk.IntVar variable for multiple_choice_options_value
            multiple_choice_options_value.set(exercise.get('MULTIPLE_CHOICE_OPTIONS_VALUE', 0))
            
            # Set the value of the global tk.IntVar variable for underscore_lines_var
            underscore_lines_var.set(exercise.get('UNDERSCORE_LINES_VAR', 0))
            
            break  # Stop searching for the exercise once found


# Assuming you have a button for adding content to the current exercise
add_to_current_button = tk.Button(catalogue_frame, text="Add to current exercise", command=copy_selected_items)
add_to_current_button.grid(row=2, column=1, pady=5)




# Extract titles from the file


exercise_entries=[]
root.mainloop()
